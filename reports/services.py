from decimal import Decimal
from datetime import datetime, timedelta
from io import BytesIO

from django.db.models import Sum, Q
from django.utils import timezone

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from store.models import Sale, SaleItem, Product, Stock, Store, StockAdjustment
from .models import CashOut, DailyNote, DailyClosing
from decimal import Decimal, ROUND_HALF_UP
from datetime import datetime, timedelta
from django.utils import timezone



def money(value):
    if value is None:
        return "0"
    try:
        return f"{Decimal(str(value)):,.0f}"
    except Exception:
        return str(value)


def calculate_stock_value(store):
    """Current inventory value = Σ (qty × buy_price)."""
    stocks = Stock.objects.filter(store=store).select_related('product')
    total = Decimal('0')
    for s in stocks:
        if s.product and s.product.buy_price:
            total += Decimal(s.quantity or 0) * Decimal(s.product.buy_price)
    return total


def get_stock_movements(store, report_date):
    """
    Auto stock in / stock out from StockAdjustment on that date.
    - increase  → stock in  (to_store)
    - decrease  → stock out (from_store)
    - transfer  → out from from_store, in to to_store
    Value = quantity × unit_price (or product.buy_price if unit_price empty).
    """
    start = timezone.make_aware(datetime.combine(report_date, datetime.min.time()))
    end = timezone.make_aware(datetime.combine(report_date, datetime.max.time()))

    adjustments = StockAdjustment.objects.filter(
        adjustment_date__range=(start, end)
    ).select_related('product', 'from_store', 'to_store')

    stock_in = Decimal('0')
    stock_out = Decimal('0')

    for adj in adjustments:
        qty = Decimal(adj.quantity or 0)
        unit = adj.unit_price
        if unit is None and adj.product_id and adj.product:
            unit = adj.product.buy_price or Decimal('0')
        unit = Decimal(unit or 0)
        value = qty * unit

        if adj.adjustment_type == 'increase':
            if adj.to_store_id == store.id:
                stock_in += value
        elif adj.adjustment_type == 'decrease':
            if adj.from_store_id == store.id:
                stock_out += value
        elif adj.adjustment_type == 'transfer':
            if adj.to_store_id == store.id:
                stock_in += value
            if adj.from_store_id == store.id:
                stock_out += value

    return stock_in, stock_out

def q(v):
    """Quantize any value to 2 decimal places. Never raises."""
    try:
        d = Decimal(str(v or 0))
        if not d.is_finite():
            return Decimal('0.00')
        return d.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    except Exception:
        return Decimal('0.00')


def get_day_data(store, report_date):
    start = timezone.make_aware(datetime.combine(report_date, datetime.min.time()))
    end = timezone.make_aware(datetime.combine(report_date, datetime.max.time()))

    sales_qs = Sale.objects.filter(store=store, sale_date__range=(start, end))
    items = (
        SaleItem.objects
        .filter(sale__in=sales_qs)
        .select_related('product', 'sale')
        .order_by('id')
    )

    lines = []
    total_product_sales = Decimal('0')
    total_buy = Decimal('0')
    total_profit = Decimal('0')

    for item in items:
        qty = item.quantity or 0
        unit_sell = item.sold_price or Decimal('0')
        line_total = item.subtotal if item.subtotal is not None else (qty * unit_sell)
        unit_buy = Decimal('0')
        if item.product_id and item.product:
            unit_buy = item.product.buy_price or Decimal('0')
        line_buy = unit_buy * qty
        line_profit = line_total - line_buy

        lines.append({
            'item_name': item.product_name,
            'qty': qty,
            'unit_price': unit_sell,
            'total_sale': line_total,
            'unit_buy': unit_buy,
            'total_buy': line_buy,
            'profit': line_profit,
        })
        total_product_sales += line_total
        total_buy += line_buy
        total_profit += line_profit

    cash_outs = list(CashOut.objects.filter(store=store, date=report_date))
    total_cash_out = sum((c.amount for c in cash_outs), Decimal('0'))
    notes = list(DailyNote.objects.filter(store=store, date=report_date))
    closing = DailyClosing.objects.filter(store=store, date=report_date).first()

    # Extra income
    advance = closing.advance_payments if closing else Decimal('0')
    maintenance = closing.maintenance_income if closing else Decimal('0')

    # Maintenance is part of daily sales (no COGS)
    total_sales = total_product_sales + maintenance
    # Profit includes maintenance as pure profit
    total_profit = total_profit + maintenance

    # Auto stock movements from adjustments
    auto_stock_in, auto_stock_out = get_stock_movements(store, report_date)
    stock_in = auto_stock_in
    stock_out = auto_stock_out
    if closing:
        if auto_stock_in == 0 and closing.stock_in:
            stock_in = closing.stock_in
        if auto_stock_out == 0 and closing.stock_out:
            stock_out = closing.stock_out

    # Auto opening / closing stock
    current_stock_value = calculate_stock_value(store)
    prev_closing = DailyClosing.objects.filter(
        store=store, date=report_date - timedelta(days=1)
    ).first()

    if prev_closing and prev_closing.closing_stock:
        opening_stock = prev_closing.closing_stock
        closing_stock = opening_stock - total_buy + stock_in - stock_out
    else:
        closing_stock = current_stock_value
        opening_stock = closing_stock + total_buy - stock_in + stock_out

    # Cash figures
    opening_balance = closing.opening_balance if closing else Decimal('0')
    bank_total = closing.bank_total if closing else Decimal('0')
    lipa_namba = closing.lipa_namba_total if closing else Decimal('0')

    cash_in_hand = (
        opening_balance
        + total_sales
        + advance
        - bank_total
        - lipa_namba
        - total_cash_out
    )
    after_sales_balance = opening_balance + total_sales + advance
    system_excel_balance = after_sales_balance - bank_total

    return {
        'date': report_date,
        'store': store,
        'lines': lines,
        'cash_outs': cash_outs,
        'notes': notes,
        'total_product_sales': q(total_product_sales),
        'total_sales': q(total_sales),
        'maintenance_income': q(maintenance),
        'advance_payments': q(advance),
        'total_buy': q(total_buy),
        'total_profit': q(total_profit),
        'total_cash_out': q(total_cash_out),
        'balance': q(total_sales - total_cash_out),
        'opening_balance': q(opening_balance),
        'after_sales_balance': q(after_sales_balance),
        'bank_total': q(bank_total),
        'system_excel_balance': q(system_excel_balance),
        'cash_in_hand': q(cash_in_hand),
        'opening_stock': q(opening_stock),
        'closing_stock': q(closing_stock),
        'stock_in': q(stock_in),
        'stock_out': q(stock_out),
        'lipa_namba': q(lipa_namba),
        'day_net_profit': q(total_profit - total_cash_out),
        'current_stock_value': q(current_stock_value),
    }


# ───────────────────────── EXCEL (English) ─────────────────────────

def generate_excel_report(store, report_date):
    data = get_day_data(store, report_date)
    wb = Workbook()
    ws = wb.active
    ws.title = report_date.strftime("%d.%m.%Y %a").upper()

    header_font = Font(bold=True, size=11)
    title_font = Font(bold=True, size=14)
    bold = Font(bold=True)
    green_fill = PatternFill("solid", fgColor="C6EFCE")
    yellow_fill = PatternFill("solid", fgColor="FFEB9C")
    grey_fill = PatternFill("solid", fgColor="D9E1F2")

    widths = {
        'A': 30, 'B': 8, 'C': 12, 'D': 12, 'E': 4,
        'F': 12, 'G': 20, 'H': 4, 'I': 14,
        'J': 4, 'K': 4, 'L': 12, 'M': 12, 'N': 4,
        'O': 12, 'P': 4, 'Q': 14, 'R': 4, 'S': 4, 'T': 14
    }
    for col, w in widths.items():
        ws.column_dimensions[col].width = w

    ws['A1'] = report_date.strftime("%d.%m.%Y %a").upper()
    ws['A1'].font = title_font

    headers = [
        ('A2', 'ITEMS'), ('B2', 'QTY'), ('C2', 'PRICE'), ('D2', 'TOTAL'),
        ('F2', 'CASH OUT'), ('G2', 'PURPOSE'), ('I2', 'BALANCE'),
        ('L2', 'BUYING'), ('M2', 'BUY TOTAL'), ('O2', 'PROFIT'),
        ('Q2', 'DAY NET PROFIT'), ('T2', 'STOCK VALUE'),
    ]
    for cell, text in headers:
        ws[cell] = text
        ws[cell].font = header_font
        ws[cell].fill = grey_fill

    row = 3
    for line in data['lines']:
        ws[f'A{row}'] = line['item_name']
        ws[f'B{row}'] = line['qty']
        ws[f'C{row}'] = float(line['unit_price'])
        ws[f'D{row}'] = float(line['total_sale'])
        ws[f'L{row}'] = float(line['unit_buy'])
        ws[f'M{row}'] = float(line['total_buy'])
        ws[f'O{row}'] = float(line['profit'])
        for col in ['C', 'D', 'L', 'M', 'O']:
            ws[f'{col}{row}'].number_format = '#,##0'
        row += 1

    # Maintenance income as a sales line
    if data['maintenance_income'] and data['maintenance_income'] > 0:
        ws[f'A{row}'] = "Device Maintenance / Repair"
        ws[f'B{row}'] = 1
        ws[f'C{row}'] = float(data['maintenance_income'])
        ws[f'D{row}'] = float(data['maintenance_income'])
        ws[f'L{row}'] = 0
        ws[f'M{row}'] = 0
        ws[f'O{row}'] = float(data['maintenance_income'])
        for col in ['C', 'D', 'O']:
            ws[f'{col}{row}'].number_format = '#,##0'
        row += 1

    cash_row = 3
    for c in data['cash_outs']:
        ws[f'F{cash_row}'] = float(c.amount)
        ws[f'G{cash_row}'] = c.purpose
        ws[f'F{cash_row}'].number_format = '#,##0'
        cash_row += 1

    total_row = max(row, cash_row) + 1
    ws[f'D{total_row}'] = float(data['total_sales'])
    ws[f'F{total_row}'] = float(data['total_cash_out'])
    ws[f'I{total_row}'] = float(data['balance'])
    ws[f'M{total_row}'] = float(data['total_buy'])
    ws[f'O{total_row}'] = float(data['total_profit'])
    ws[f'Q{total_row}'] = float(data['day_net_profit'])
    for col in ['D', 'F', 'I', 'M', 'O', 'Q']:
        ws[f'{col}{total_row}'].number_format = '#,##0'
        ws[f'{col}{total_row}'].font = bold
        ws[f'{col}{total_row}'].fill = green_fill

    r = total_row + 2
    prev = report_date - timedelta(days=1)
    ws[f'A{r}'] = f"Previous balance ({prev.strftime('%d.%m.%Y')})"
    ws[f'I{r}'] = float(data['opening_balance'])
    ws[f'I{r}'].number_format = '#,##0'

    r += 2
    ws[f'A{r}'] = "After sales balance"
    ws[f'I{r}'] = float(data['after_sales_balance'])
    ws[f'I{r}'].number_format = '#,##0'
    ws[f'I{r}'].font = bold

    r += 2
    ws[f'A{r}'] = "TOTAL BANK DEPOSIT"
    ws[f'I{r}'] = float(data['bank_total'])
    ws[f'I{r}'].number_format = '#,##0'

    r += 2
    ws[f'A{r}'] = "SYSTEM & EXCEL BALANCED"
    ws[f'I{r}'] = float(data['system_excel_balance'])
    ws[f'I{r}'].number_format = '#,##0'
    ws[f'I{r}'].fill = yellow_fill

    r += 2
    ws[f'A{r}'] = f"Cash on hand: TSH {money(data['cash_in_hand'])}/="

    r += 2
    if data['advance_payments']:
        ws[f'A{r}'] = f"Advance payments received: TSH {money(data['advance_payments'])}/="
        r += 1
    if data['maintenance_income']:
        ws[f'A{r}'] = f"Device maintenance income: TSH {money(data['maintenance_income'])}/="
        r += 1

    for note in data['notes']:
        ws[f'A{r}'] = note.note
        r += 1

    if data['lipa_namba']:
        ws[f'A{r}'] = f"TOTAL MOBILE MONEY (LIPA NAMBA): TSH {money(data['lipa_namba'])}/="
        r += 1

    r += 1
    ws[f'A{r}'] = "SYSTEM & EXCEL BALANCED"
    r += 1
    ws[f'A{r}'] = f"OPENING STOCK: {money(data['opening_stock'])}/="
    r += 1
    ws[f'A{r}'] = f"STOCK RECEIVED (IN): {money(data['stock_in'])}/="
    r += 1
    ws[f'A{r}'] = f"STOCK ISSUED (OUT): {money(data['stock_out'])}/="
    r += 1
    ws[f'A{r}'] = f"CLOSING STOCK: {money(data['closing_stock'])}/="
    r += 1
    ws[f'A{r}'] = f"TODAY'S SALES: {money(data['total_sales'])}/="
    ws[f'A{r}'].font = bold
    r += 1
    ws[f'A{r}'] = f"CLOSING CASH POSITION: TSH {money(data['after_sales_balance'])}/="
    r += 1
    ws[f'A{r}'] = f"CASH IN HAND: TSH {money(data['cash_in_hand'])}/="
    ws[f'A{r}'].font = bold
    ws[f'A{r}'].fill = green_fill

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


# ───────────────────────── PDF (English) ─────────────────────────

def generate_pdf_report(store, report_date):
    data = get_day_data(store, report_date)
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=landscape(A4),
        leftMargin=12*mm, rightMargin=12*mm, topMargin=12*mm, bottomMargin=12*mm
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('T', parent=styles['Heading1'], fontSize=14, spaceAfter=8)
    normal = styles['Normal']
    story = []

    story.append(Paragraph(
        f"Daily Sales Report — {store.name} — {report_date.strftime('%d/%m/%Y (%A)')}",
        title_style
    ))
    story.append(Spacer(1, 6))

    table_data = [['ITEMS', 'QTY', 'PRICE', 'TOTAL', 'BUYING', 'BUY TOTAL', 'PROFIT']]
    for line in data['lines']:
        table_data.append([
            line['item_name'], str(line['qty']),
            money(line['unit_price']), money(line['total_sale']),
            money(line['unit_buy']), money(line['total_buy']), money(line['profit']),
        ])
    if data['maintenance_income'] and data['maintenance_income'] > 0:
        table_data.append([
            "Device Maintenance / Repair", "1",
            money(data['maintenance_income']), money(data['maintenance_income']),
            "0", "0", money(data['maintenance_income']),
        ])
    table_data.append([
        'TOTAL', '', '',
        money(data['total_sales']), '',
        money(data['total_buy']), money(data['total_profit']),
    ])

    t = Table(table_data, colWidths=[70*mm, 15*mm, 25*mm, 28*mm, 25*mm, 28*mm, 28*mm])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#647688')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('ALIGN', (1, 0), (-1, -1), 'RIGHT'),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.grey),
        ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#C6EFCE')),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(t)
    story.append(Spacer(1, 10))

    if data['cash_outs']:
        story.append(Paragraph("<b>Daily Expenses (Cash Out)</b>", normal))
        co_data = [['Amount', 'Purpose']]
        for c in data['cash_outs']:
            co_data.append([money(c.amount), c.purpose])
        co_data.append([money(data['total_cash_out']), 'TOTAL'])
        t2 = Table(co_data, colWidths=[30*mm, 80*mm])
        t2.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#647688')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.4, colors.grey),
            ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#FFEB9C')),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ]))
        story.append(t2)
        story.append(Spacer(1, 8))

    summary = [
        f"<b>Product Sales:</b> {money(data['total_product_sales'])}/=",
        f"<b>Device Maintenance Income:</b> {money(data['maintenance_income'])}/=",
        f"<b>Total Sales (Today):</b> {money(data['total_sales'])}/=",
        f"<b>Advance Payments Received:</b> {money(data['advance_payments'])}/=",
        f"<b>Total Buying (COGS):</b> {money(data['total_buy'])}/=",
        f"<b>Day Gross Profit:</b> {money(data['total_profit'])}/=",
        f"<b>Daily Expenses:</b> {money(data['total_cash_out'])}/=",
        f"<b>Opening Balance:</b> {money(data['opening_balance'])}/=",
        f"<b>Bank Deposit:</b> {money(data['bank_total'])}/=",
        f"<b>Mobile Money (Lipa Namba):</b> {money(data['lipa_namba'])}/=",
        f"<b>Cash in Hand:</b> {money(data['cash_in_hand'])}/=",
        f"<b>Opening Stock:</b> {money(data['opening_stock'])}/=",
        f"<b>Stock Received (In):</b> {money(data['stock_in'])}/=",
        f"<b>Stock Issued (Out):</b> {money(data['stock_out'])}/=",
        f"<b>Closing Stock:</b> {money(data['closing_stock'])}/=",
    ]
    for line in summary:
        story.append(Paragraph(line, normal))
        story.append(Spacer(1, 2))

    if data['notes']:
        story.append(Spacer(1, 6))
        story.append(Paragraph("<b>Comments</b>", normal))
        for n in data['notes']:
            story.append(Paragraph(f"• {n.note}", normal))

    doc.build(story)
    buffer.seek(0)
    return buffer


# ───────────────────────── WORD (English) ─────────────────────────

def generate_word_report(store, report_date):
    data = get_day_data(store, report_date)
    doc = Document()

    title = doc.add_heading(f"Daily Sales Report — {store.name}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(report_date.strftime("%d/%m/%Y (%A)"))
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Sales", level=2)
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['ITEMS', 'QTY', 'PRICE', 'TOTAL', 'BUYING', 'BUY TOTAL', 'PROFIT']):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True

    for line in data['lines']:
        row = table.add_row().cells
        row[0].text = line['item_name']
        row[1].text = str(line['qty'])
        row[2].text = money(line['unit_price'])
        row[3].text = money(line['total_sale'])
        row[4].text = money(line['unit_buy'])
        row[5].text = money(line['total_buy'])
        row[6].text = money(line['profit'])

    if data['maintenance_income'] and data['maintenance_income'] > 0:
        row = table.add_row().cells
        row[0].text = "Device Maintenance / Repair"
        row[1].text = "1"
        row[2].text = money(data['maintenance_income'])
        row[3].text = money(data['maintenance_income'])
        row[4].text = "0"
        row[5].text = "0"
        row[6].text = money(data['maintenance_income'])

    row = table.add_row().cells
    row[0].text = "TOTAL"
    row[3].text = money(data['total_sales'])
    row[5].text = money(data['total_buy'])
    row[6].text = money(data['total_profit'])
    for cell in row:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    if data['cash_outs']:
        doc.add_heading("Daily Expenses", level=2)
        co = doc.add_table(rows=1, cols=2)
        co.style = 'Table Grid'
        co.rows[0].cells[0].text = "Amount"
        co.rows[0].cells[1].text = "Purpose"
        for c in data['cash_outs']:
            r = co.add_row().cells
            r[0].text = money(c.amount)
            r[1].text = c.purpose
        r = co.add_row().cells
        r[0].text = money(data['total_cash_out'])
        r[1].text = "TOTAL"

    doc.add_heading("Day Summary", level=2)
    lines = [
        f"Product Sales: {money(data['total_product_sales'])}/=",
        f"Device Maintenance Income: {money(data['maintenance_income'])}/=",
        f"Total Sales (Today): {money(data['total_sales'])}/=",
        f"Advance Payments Received: {money(data['advance_payments'])}/=",
        f"Total Buying (COGS): {money(data['total_buy'])}/=",
        f"Day Gross Profit: {money(data['total_profit'])}/=",
        f"Daily Expenses: {money(data['total_cash_out'])}/=",
        f"Opening Balance: {money(data['opening_balance'])}/=",
        f"Bank Deposit: {money(data['bank_total'])}/=",
        f"Mobile Money (Lipa Namba): {money(data['lipa_namba'])}/=",
        f"Cash in Hand: {money(data['cash_in_hand'])}/=",
        f"Opening Stock: {money(data['opening_stock'])}/=",
        f"Stock Received (In): {money(data['stock_in'])}/=",
        f"Stock Issued (Out): {money(data['stock_out'])}/=",
        f"Closing Stock: {money(data['closing_stock'])}/=",
    ]
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)

    if data['notes']:
        doc.add_heading("Comments", level=2)
        for n in data['notes']:
            doc.add_paragraph(n.note, style='List Bullet')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


from datetime import date as date_cls
from calendar import monthrange


def resolve_period(period, start_str=None, end_str=None, ref_date=None):
    """
    period: daily | weekly | monthly | yearly | custom
    Returns (start_date, end_date)
    """
    today = ref_date or date_cls.today()

    if period == 'daily':
        return today, today

    if period == 'weekly':
        # Monday–Sunday of current week
        start = today - timedelta(days=today.weekday())
        end = start + timedelta(days=6)
        return start, end

    if period == 'monthly':
        start = today.replace(day=1)
        last = monthrange(today.year, today.month)[1]
        end = today.replace(day=last)
        return start, end

    if period == 'yearly':
        start = today.replace(month=1, day=1)
        end = today.replace(month=12, day=31)
        return start, end

    if period == 'custom' and start_str and end_str:
        try:
            start = datetime.strptime(start_str, '%Y-%m-%d').date()
            end = datetime.strptime(end_str, '%Y-%m-%d').date()
            if start > end:
                start, end = end, start
            return start, end
        except ValueError:
            pass

    # fallback: today
    return today, today


def get_expenses_report(store, start_date, end_date):
    qs = CashOut.objects.filter(
        store=store,
        date__gte=start_date,
        date__lte=end_date,
    ).order_by('date', 'id')

    total = sum((c.amount for c in qs), Decimal('0'))
    # Group by date for summary
    by_date = {}
    for c in qs:
        by_date.setdefault(c.date, Decimal('0'))
        by_date[c.date] += c.amount

    return {
        'items': list(qs),
        'total': total,
        'by_date': sorted(by_date.items()),
        'count': qs.count(),
        'start': start_date,
        'end': end_date,
    }


def get_stock_adjustments_report(store, start_date, end_date):
    start_dt = timezone.make_aware(datetime.combine(start_date, datetime.min.time()))
    end_dt = timezone.make_aware(datetime.combine(end_date, datetime.max.time()))

    qs = StockAdjustment.objects.filter(
        adjustment_date__range=(start_dt, end_dt)
    ).filter(
        Q(from_store=store) | Q(to_store=store)
    ).select_related('product', 'from_store', 'to_store', 'adjusted_by').order_by('-adjustment_date')

    total_in = Decimal('0')
    total_out = Decimal('0')
    rows = []

    for adj in qs:
        qty = Decimal(adj.quantity or 0)
        unit = adj.unit_price
        if unit is None and adj.product_id and adj.product:
            unit = adj.product.buy_price or Decimal('0')
        unit = Decimal(unit or 0)
        value = qty * unit

        direction = ''
        if adj.adjustment_type == 'increase' and adj.to_store_id == store.id:
            direction = 'IN'
            total_in += value
        elif adj.adjustment_type == 'decrease' and adj.from_store_id == store.id:
            direction = 'OUT'
            total_out += value
        elif adj.adjustment_type == 'transfer':
            if adj.to_store_id == store.id:
                direction = 'IN (transfer)'
                total_in += value
            if adj.from_store_id == store.id:
                direction = 'OUT (transfer)' if direction == '' else 'IN+OUT'
                total_out += value

        rows.append({
            'adj': adj,
            'qty': qty,
            'unit': unit,
            'value': value,
            'direction': direction,
        })

    return {
        'rows': rows,
        'total_in': total_in,
        'total_out': total_out,
        'count': len(rows),
        'start': start_date,
        'end': end_date,
    }


def get_sales_period_report(store, start_date, end_date):
    start_dt = timezone.make_aware(datetime.combine(start_date, datetime.min.time()))
    end_dt = timezone.make_aware(datetime.combine(end_date, datetime.max.time()))

    sales = Sale.objects.filter(
        store=store,
        sale_date__range=(start_dt, end_dt),
    ).order_by('-sale_date')

    items = (
        SaleItem.objects
        .filter(sale__in=sales)
        .select_related('product', 'sale')
        .order_by('-sale__sale_date', 'id')
    )

    lines = []
    total_sales = Decimal('0')
    total_buy = Decimal('0')
    total_profit = Decimal('0')
    total_qty = 0

    for item in items:
        qty = item.quantity or 0
        unit_sell = item.sold_price or Decimal('0')
        line_total = item.subtotal if item.subtotal is not None else (qty * unit_sell)
        unit_buy = Decimal('0')
        if item.product_id and item.product:
            unit_buy = item.product.buy_price or Decimal('0')
        line_buy = unit_buy * qty
        profit = line_total - line_buy

        lines.append({
            'date': item.sale.sale_date.date() if item.sale else None,
            'sale_id': item.sale_id,
            'item_name': item.product_name,
            'qty': qty,
            'unit_price': unit_sell,
            'total_sale': line_total,
            'unit_buy': unit_buy,
            'total_buy': line_buy,
            'profit': profit,
        })
        total_sales += line_total
        total_buy += line_buy
        total_profit += profit
        total_qty += qty

    # Maintenance income in period (from DailyClosing)
    closings = DailyClosing.objects.filter(
        store=store, date__gte=start_date, date__lte=end_date
    )
    maintenance = sum((c.maintenance_income or Decimal('0') for c in closings), Decimal('0'))
    advances = sum((c.advance_payments or Decimal('0') for c in closings), Decimal('0'))

    return {
        'lines': lines,
        'sales_count': sales.count(),
        'items_count': len(lines),
        'total_qty': total_qty,
        'total_sales': total_sales,
        'total_buy': total_buy,
        'total_profit': total_profit + maintenance,
        'maintenance': maintenance,
        'advances': advances,
        'grand_sales': total_sales + maintenance,
        'start': start_date,
        'end': end_date,
    }


# ── Simple Excel exports for period reports ──

def export_expenses_excel(store, start_date, end_date):
    data = get_expenses_report(store, start_date, end_date)
    wb = Workbook()
    ws = wb.active
    ws.title = "Expenses"
    ws.append(["Date", "Amount (TZS)", "Purpose"])
    for c in data['items']:
        ws.append([c.date.strftime('%Y-%m-%d'), float(c.amount), c.purpose])
    ws.append([])
    ws.append(["TOTAL", float(data['total']), ""])
    for col in ['A', 'B', 'C']:
        ws.column_dimensions[col].width = 20 if col != 'C' else 40
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def export_stock_excel(store, start_date, end_date):
    data = get_stock_adjustments_report(store, start_date, end_date)
    wb = Workbook()
    ws = wb.active
    ws.title = "Stock Adjustments"
    ws.append(["Date", "Product", "Type", "Direction", "Qty", "Unit Price", "Value", "From", "To", "Reason", "By"])
    for r in data['rows']:
        adj = r['adj']
        ws.append([
            adj.adjustment_date.strftime('%Y-%m-%d %H:%M'),
            str(adj.product) if adj.product else '',
            adj.get_adjustment_type_display(),
            r['direction'],
            float(r['qty']),
            float(r['unit']),
            float(r['value']),
            adj.from_store.name if adj.from_store else '',
            adj.to_store.name if adj.to_store else '',
            adj.reason or '',
            str(adj.adjusted_by) if adj.adjusted_by else '',
        ])
    ws.append([])
    ws.append(["", "", "", "TOTAL IN", "", "", float(data['total_in'])])
    ws.append(["", "", "", "TOTAL OUT", "", "", float(data['total_out'])])
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def export_sales_excel(store, start_date, end_date):
    data = get_sales_period_report(store, start_date, end_date)
    wb = Workbook()
    ws = wb.active
    ws.title = "Sales"
    ws.append(["Date", "Sale #", "Item", "Qty", "Unit Price", "Total", "Buy Unit", "Buy Total", "Profit"])
    for line in data['lines']:
        ws.append([
            line['date'].strftime('%Y-%m-%d') if line['date'] else '',
            line['sale_id'],
            line['item_name'],
            line['qty'],
            float(line['unit_price']),
            float(line['total_sale']),
            float(line['unit_buy']),
            float(line['total_buy']),
            float(line['profit']),
        ])
    ws.append([])
    ws.append(["", "", "TOTAL", data['total_qty'], "",
               float(data['total_sales']), "", float(data['total_buy']), float(data['total_profit'])])
    if data['maintenance']:
        ws.append(["", "", "Maintenance income", "", "", float(data['maintenance'])])
        ws.append(["", "", "GRAND SALES", "", "", float(data['grand_sales'])])
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer


def _pdf_table_style(header=True):
    style = [
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]
    if header:
        style += [
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#647688')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ]
    return TableStyle(style)


def export_expenses_pdf(store, start_date, end_date):
    data = get_expenses_report(store, start_date, end_date)
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=15*mm, rightMargin=15*mm, topMargin=15*mm, bottomMargin=15*mm
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('T', parent=styles['Heading1'], fontSize=13, spaceAfter=6)
    story = []

    story.append(Paragraph(
        f"Expenses Report — {store.name}<br/>"
        f"{start_date.strftime('%d %b %Y')} to {end_date.strftime('%d %b %Y')}",
        title_style
    ))
    story.append(Spacer(1, 8))

    table_data = [['#', 'Date', 'Amount (TZS)', 'Purpose']]
    for i, c in enumerate(data['items'], 1):
        table_data.append([
            str(i),
            c.date.strftime('%d %b %Y'),
            money(c.amount),
            c.purpose,
        ])
    table_data.append(['', 'TOTAL', money(data['total']), ''])

    t = Table(table_data, colWidths=[15*mm, 35*mm, 40*mm, 80*mm])
    style = _pdf_table_style()
    style.add('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#C6EFCE'))
    style.add('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold')
    style.add('ALIGN', (2, 1), (2, -1), 'RIGHT')
    t.setStyle(style)
    story.append(t)

    story.append(Spacer(1, 12))
    story.append(Paragraph(f"<b>Total entries:</b> {data['count']}", styles['Normal']))
    story.append(Paragraph(f"<b>Total expenses:</b> {money(data['total'])} TZS", styles['Normal']))

    doc.build(story)
    buffer.seek(0)
    return buffer


def export_stock_pdf(store, start_date, end_date):
    data = get_stock_adjustments_report(store, start_date, end_date)
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=landscape(A4),
        leftMargin=12*mm, rightMargin=12*mm, topMargin=12*mm, bottomMargin=12*mm
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('T', parent=styles['Heading1'], fontSize=13, spaceAfter=6)
    story = []

    story.append(Paragraph(
        f"Stock Adjustments Report — {store.name}<br/>"
        f"{start_date.strftime('%d %b %Y')} to {end_date.strftime('%d %b %Y')}",
        title_style
    ))
    story.append(Spacer(1, 8))

    table_data = [['Date', 'Product', 'Type', 'Dir', 'Qty', 'Unit', 'Value', 'From', 'To', 'Reason']]
    for r in data['rows']:
        adj = r['adj']
        table_data.append([
            adj.adjustment_date.strftime('%d/%m/%y %H:%M'),
            str(adj.product)[:22] if adj.product else '',
            adj.get_adjustment_type_display()[:12],
            r['direction'],
            money(r['qty']),
            money(r['unit']),
            money(r['value']),
            (adj.from_store.name[:12] if adj.from_store else '—'),
            (adj.to_store.name[:12] if adj.to_store else '—'),
            (adj.reason or '')[:20],
        ])

    t = Table(table_data, colWidths=[28*mm, 40*mm, 22*mm, 22*mm, 18*mm, 22*mm, 25*mm, 25*mm, 25*mm, 30*mm])
    style = _pdf_table_style()
    style.add('ALIGN', (4, 1), (6, -1), 'RIGHT')
    t.setStyle(style)
    story.append(t)

    story.append(Spacer(1, 12))
    story.append(Paragraph(f"<b>Stock In:</b> {money(data['total_in'])} TZS", styles['Normal']))
    story.append(Paragraph(f"<b>Stock Out:</b> {money(data['total_out'])} TZS", styles['Normal']))
    story.append(Paragraph(f"<b>Movements:</b> {data['count']}", styles['Normal']))

    doc.build(story)
    buffer.seek(0)
    return buffer


def export_sales_pdf(store, start_date, end_date):
    data = get_sales_period_report(store, start_date, end_date)
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=landscape(A4),
        leftMargin=12*mm, rightMargin=12*mm, topMargin=12*mm, bottomMargin=12*mm
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('T', parent=styles['Heading1'], fontSize=13, spaceAfter=6)
    story = []

    story.append(Paragraph(
        f"Sales Report — {store.name}<br/>"
        f"{start_date.strftime('%d %b %Y')} to {end_date.strftime('%d %b %Y')}",
        title_style
    ))
    story.append(Spacer(1, 8))

    table_data = [['Date', 'Sale #', 'Item', 'Qty', 'Price', 'Total', 'Buy Unit', 'Buy Total', 'Profit']]
    for line in data['lines']:
        table_data.append([
            line['date'].strftime('%d/%m/%y') if line['date'] else '',
            str(line['sale_id']),
            (line['item_name'] or '')[:28],
            str(line['qty']),
            money(line['unit_price']),
            money(line['total_sale']),
            money(line['unit_buy']),
            money(line['total_buy']),
            money(line['profit']),
        ])
    table_data.append([
        '', '', 'TOTAL', str(data['total_qty']), '',
        money(data['total_sales']), '',
        money(data['total_buy']), money(data['total_profit']),
    ])

    t = Table(table_data, colWidths=[22*mm, 18*mm, 55*mm, 15*mm, 25*mm, 28*mm, 25*mm, 28*mm, 28*mm])
    style = _pdf_table_style()
    style.add('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#C6EFCE'))
    style.add('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold')
    style.add('ALIGN', (3, 1), (-1, -1), 'RIGHT')
    t.setStyle(style)
    story.append(t)

    story.append(Spacer(1, 12))
    story.append(Paragraph(f"<b>Product sales:</b> {money(data['total_sales'])} TZS", styles['Normal']))
    if data['maintenance']:
        story.append(Paragraph(f"<b>Maintenance income:</b> {money(data['maintenance'])} TZS", styles['Normal']))
        story.append(Paragraph(f"<b>Grand sales:</b> {money(data['grand_sales'])} TZS", styles['Normal']))
    story.append(Paragraph(f"<b>Gross profit:</b> {money(data['total_profit'])} TZS", styles['Normal']))
    story.append(Paragraph(f"<b>Sales count:</b> {data['sales_count']} &nbsp;|&nbsp; <b>Items:</b> {data['total_qty']}", styles['Normal']))

    doc.build(story)
    buffer.seek(0)
    return buffer